#!/usr/bin/env python3
"""
生成包含多种颜色格式的复杂 DOCX 模板
确保颜色是直接在运行级别设置的，而不是通过样式引用
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 1. 标题 - 包含多种颜色
title = doc.add_heading('复杂格式测试文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)  # 深灰色

# 2. 副标题 - 红色
subtitle = doc.add_heading('格式保留测试', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色

# 3. 段落 - 包含多种颜色的文字
para1 = doc.add_paragraph()
para1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 添加不同颜色的文字
run1 = para1.add_run('这是')
run1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

run2 = para1.add_run('蓝色')
run2.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 蓝色
run2.font.bold = True

run3 = para1.add_run('的文字，')
run3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

run4 = para1.add_run('绿色')
run4.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)  # 绿色
run4.font.italic = True

run5 = para1.add_run('的文字，')
run5.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

run6 = para1.add_run('红色')
run6.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
run6.font.bold = True
run6.font.italic = True

run7 = para1.add_run('的文字。')
run7.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

# 4. 段落 - 包含单个字符的不同颜色
para2 = doc.add_paragraph()
para2.alignment = WD_ALIGN_PARAGRAPH.LEFT

text = "彩虹文字："
for i, char in enumerate(text):
    run = para2.add_run(char)
    if i < len(text) - 1:  # 最后一个字符保持默认颜色
        # 使用彩虹色
        colors = [
            RGBColor(0xFF, 0x00, 0x00),  # 红
            RGBColor(0xFF, 0x7F, 0x00),  # 橙
            RGBColor(0xFF, 0xFF, 0x00),  # 黄
            RGBColor(0x00, 0xFF, 0x00),  # 绿
            RGBColor(0x00, 0x00, 0xFF),  # 蓝
            RGBColor(0x4B, 0x00, 0x82),  # 靛
            RGBColor(0x94, 0x00, 0xD3),  # 紫
        ]
        run.font.color.rgb = colors[i % len(colors)]

# 5. 段落 - 包含标点符号的不同颜色
para3 = doc.add_paragraph()
para3.alignment = WD_ALIGN_PARAGRAPH.LEFT

run_word1 = para3.add_run('单词')
run_word1.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 蓝色

run_punct1 = para3.add_run('，')
run_punct1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色

run_word2 = para3.add_run('标点')
run_word2.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)  # 绿色

run_punct2 = para3.add_run('。')
run_punct2.font.color.rgb = RGBColor(0xFF, 0x00, 0xFF)  # 紫色

# 6. 段落 - 包含下划线和颜色的组合
para4 = doc.add_paragraph()
para4.alignment = WD_ALIGN_PARAGRAPH.LEFT

run_under1 = para4.add_run('下划线')
run_under1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
run_under1.font.underline = True

run_normal = para4.add_run(' 普通文字 ')

run_under2 = para4.add_run('下划线')
run_under2.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 蓝色
run_under2.font.underline = True

# 7. 段落 - 包含不同字号的彩色文字
para5 = doc.add_paragraph()
para5.alignment = WD_ALIGN_PARAGRAPH.LEFT

run_size1 = para5.add_run('小号')
run_size1.font.size = Pt(10)
run_size1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色

run_normal2 = para5.add_run(' 正常 ')

run_size2 = para5.add_run('大号')
run_size2.font.size = Pt(16)
run_size2.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 蓝色

# 8. 段落 - 居中对齐的彩色文字
para6 = doc.add_paragraph()
para6.alignment = WD_ALIGN_PARAGRAPH.CENTER

run_center1 = para6.add_run('居中')
run_center1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
run_center1.font.bold = True

run_center2 = para6.add_run(' 文字 ')
run_center2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

run_center3 = para6.add_run('测试')
run_center3.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)  # 绿色
run_center3.font.italic = True

# 9. 段落 - 包含混合格式的复杂文本
para7 = doc.add_paragraph()
para7.alignment = WD_ALIGN_PARAGRAPH.LEFT

complex_texts = [
    ('这是', RGBColor(0x00, 0x00, 0x00), False, False),
    ('一个', RGBColor(0xFF, 0x00, 0x00), True, False),
    ('复杂', RGBColor(0x00, 0xFF, 0x00), False, True),
    ('的', RGBColor(0x00, 0x00, 0x00), False, False),
    ('格式', RGBColor(0x00, 0x00, 0xFF), True, True),
    ('测试', RGBColor(0xFF, 0x00, 0xFF), False, False),
    ('段落', RGBColor(0xFF, 0x7F, 0x00), True, True),
    ('。', RGBColor(0x00, 0x00, 0x00), False, False),
]

for text, color, bold, italic in complex_texts:
    run = para7.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic

# 10. 段落 - 包含特殊字符和颜色
para8 = doc.add_paragraph()
para8.alignment = WD_ALIGN_PARAGRAPH.LEFT

special_chars = [
    ('特殊', RGBColor(0xFF, 0x00, 0x00)),
    ('字符', RGBColor(0x00, 0xFF, 0x00)),
    ('：', RGBColor(0x00, 0x00, 0xFF)),
    ('@', RGBColor(0xFF, 0x00, 0xFF)),
    ('#', RGBColor(0xFF, 0x7F, 0x00)),
    ('$', RGBColor(0x00, 0xFF, 0xFF)),
    ('%', RGBColor(0xFF, 0xFF, 0x00)),
    ('&', RGBColor(0x00, 0x00, 0xFF)),
    ('*', RGBColor(0xFF, 0x00, 0x00)),
]

for text, color in special_chars:
    run = para8.add_run(text)
    run.font.color.rgb = color

# 保存文档
output_path = 'complex_reference.docx'
doc.save(output_path)
print(f"✅ 已生成复杂模板: {output_path}")
print(f"📄 包含 {len(doc.paragraphs)} 个段落")
print(f"🎨 包含多种颜色格式：红色、蓝色、绿色、紫色、橙色、黄色等")

